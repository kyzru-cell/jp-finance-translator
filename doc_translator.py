"""
doc_translator.py
支持 Word (.docx) 和 PDF 输入，翻译后输出双语对照 Word 文档。

输出格式：
  每段落：
    【中文】原文段落
    【日語】译文段落
    （空行）
"""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import re


@dataclass
class Paragraph:
    text: str
    style: str = "Normal"       # Word 样式名
    is_heading: bool = False
    heading_level: int = 0


@dataclass
class DocResult:
    source_path: Path
    output_path: Path
    total_paragraphs: int
    translated: int
    skipped: int                # 空段落、图片标题等
    avg_score: float


# ── 读取 ─────────────────────────────────────────────────────

def read_docx(path: Path) -> list[Paragraph]:
    from docx import Document
    doc = Document(str(path))
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_heading = p.style.name.startswith("Heading")
        level = 0
        if is_heading:
            m = re.search(r"(\d+)", p.style.name)
            level = int(m.group(1)) if m else 1
        paragraphs.append(Paragraph(
            text=text,
            style=p.style.name,
            is_heading=is_heading,
            heading_level=level,
        ))
    return paragraphs


def read_pdf(path: Path) -> list[Paragraph]:
    import pdfplumber
    paragraphs = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                # 简单启发：全大写短句可能是标题
                is_heading = len(line) < 60 and line.isupper()
                paragraphs.append(Paragraph(
                    text=line,
                    is_heading=is_heading,
                    heading_level=1 if is_heading else 0,
                ))
    return paragraphs


# ── 写出双语 Word ─────────────────────────────────────────────

def write_bilingual_docx(
    paragraphs: list[Paragraph],
    translations: list[str],
    output_path: Path,
    source_path: Path,
):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title = doc.add_heading(f"翻訳版：{source_path.stem}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para, ja_text in zip(paragraphs, translations):
        if para.is_heading:
            # 标题：中日各一行
            h = doc.add_heading(para.text, level=para.heading_level)
            h_ja = doc.add_heading(ja_text, level=para.heading_level)
            # 日语标题用灰色区分
            for run in h_ja.runs:
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            doc.add_paragraph()
        else:
            # 正文：中文段落
            p_zh = doc.add_paragraph()
            run_label = p_zh.add_run("【中文】")
            run_label.font.bold = True
            run_label.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
            run_label.font.size = Pt(9)
            p_zh.add_run(para.text)

            # 日语段落
            p_ja = doc.add_paragraph()
            run_label_ja = p_ja.add_run("【日語】")
            run_label_ja.font.bold = True
            run_label_ja.font.color.rgb = RGBColor(0x19, 0x6E, 0x28)
            run_label_ja.font.size = Pt(9)
            p_ja.add_run(ja_text)

            doc.add_paragraph()  # 空行分隔

    doc.save(str(output_path))


# ── 主流程 ────────────────────────────────────────────────────

def translate_doc(
    input_path: Path,
    output_path: Optional[Path] = None,
    glossary_id: Optional[str] = None,
    style: str = "formal",
    batch_size: int = 5,
) -> DocResult:
    """
    翻译 Word 或 PDF 文档，输出双语对照 Word 文件。

    batch_size: 每批合并翻译的段落数（减少 API 调用次数）
    """
    import orchestrator
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn

    console = Console(force_terminal=True, legacy_windows=False)

    # 确定输出路径
    if output_path is None:
        output_path = input_path.with_stem(input_path.stem + "_双語").with_suffix(".docx")

    # 读取文档
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        paragraphs = read_docx(input_path)
    elif suffix == ".pdf":
        paragraphs = read_pdf(input_path)
    else:
        raise ValueError(f"不支持的格式：{suffix}（支持 .docx / .pdf）")

    console.print(f"[cyan]读取完成[/cyan] {len(paragraphs)} 个段落")

    # 分批翻译
    translations = []
    scores = []
    skipped = 0

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console,
    ) as progress:
        task = progress.add_task("翻译中...", total=len(paragraphs))

        for i in range(0, len(paragraphs), batch_size):
            batch = paragraphs[i:i + batch_size]

            for para in batch:
                # 极短文本（纯数字/标点/1-2字）跳过 API，直接复制
                if len(para.text) <= 2 or re.fullmatch(r"[\d\s\W]+", para.text):
                    translations.append(para.text)
                    skipped += 1
                else:
                    result = orchestrator.translate(para.text, glossary_id=glossary_id, style=style)
                    translations.append(result.final)
                    scores.append(result.score)

                progress.advance(task)

    # 写出双语文档
    write_bilingual_docx(paragraphs, translations, output_path, input_path)

    avg_score = sum(scores) / len(scores) if scores else 0
    return DocResult(
        source_path=input_path,
        output_path=output_path,
        total_paragraphs=len(paragraphs),
        translated=len(paragraphs) - skipped,
        skipped=skipped,
        avg_score=round(avg_score, 1),
    )
