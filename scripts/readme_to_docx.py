"""将 README.md 转换为 Word 文档"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

README = Path("C:/Users/icoxia/.claude/skills/lark-cn2jp-finance/README.md")
OUTPUT = Path("C:/Users/icoxia/Desktop/中日金融专业翻译Skill使用说明.docx")

doc = Document()

# 页边距
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# 样式辅助
def set_heading(para, level):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: RGBColor(0x1F, 0x5C, 0x99), 2: RGBColor(0x2E, 0x75, 0xB6), 3: RGBColor(0x40, 0x40, 0x40)}
    for run in para.runs:
        run.font.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

def add_table_from_md(doc, lines):
    """解析 markdown 表格并添加到文档"""
    rows = []
    for line in lines:
        if re.match(r"^\|[-| ]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    doc.add_paragraph()

lines = README.read_text(encoding="utf-8").splitlines()
i = 0
in_code = False
code_lines = []
table_lines = []
in_table = False

while i < len(lines):
    line = lines[i]

    # 代码块
    if line.strip().startswith("```"):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            p.paragraph_format.left_indent = Inches(0.3)
            # 灰色背景效果（段落底纹）
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            doc.add_paragraph()
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # 表格
    if line.startswith("|"):
        table_lines.append(line)
        i += 1
        continue
    else:
        if table_lines:
            add_table_from_md(doc, table_lines)
            table_lines = []

    # 标题
    if line.startswith("# ") and not line.startswith("## "):
        p = doc.add_heading(line[2:].strip(), level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_heading(p, 1)
    elif line.startswith("## "):
        p = doc.add_heading(line[3:].strip(), level=2)
        set_heading(p, 2)
    elif line.startswith("### "):
        p = doc.add_heading(line[4:].strip(), level=3)
        set_heading(p, 3)
    elif line.startswith("> "):
        p = doc.add_paragraph(line[2:].strip())
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(10)
    elif line.startswith("---"):
        doc.add_paragraph()
    elif line.strip() == "":
        pass  # 空行不加段落，避免过多间距
    else:
        # 处理 **bold** 和 `code`
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        # 简单处理加粗和行内代码
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(11)
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
            else:
                if part.strip():
                    run = p.add_run(part)
                    run.font.size = Pt(11)

    i += 1

# 处理末尾可能残留的表格
if table_lines:
    add_table_from_md(doc, table_lines)

doc.save(str(OUTPUT))
print(f"已生成：{OUTPUT}")
